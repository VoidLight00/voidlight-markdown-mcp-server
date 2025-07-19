"""
DOCX file converter
"""

from typing import BinaryIO, Dict, Any, Optional, List
import logging

from ..core.base_converter import DocumentConverter, DocumentConverterResult
from ..core.stream_info import StreamInfo
from ..core.exceptions import MissingDependencyException, FileConversionException
from ..utils.format_utils import normalize_whitespace

logger = logging.getLogger(__name__)


class DocxConverter(DocumentConverter):
    """Microsoft Word DOCX file converter"""
    
    def __init__(self, korean_support: bool = True):
        super().__init__(korean_support=korean_support)
        self.priority = 0.1
        
        self.supported_extensions = ['.docx']
        self.supported_mimetypes = [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        ]
        self.category = 'documents'
        
        # Check dependencies
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check for available dependencies"""
        try:
            from docx import Document
            self.python_docx_available = True
        except ImportError:
            self.python_docx_available = False
    
    def accepts(self, file_stream: BinaryIO, stream_info: StreamInfo, **kwargs) -> bool:
        """Check if this converter can handle the file"""
        return (stream_info.matches_mimetype([
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        ]) or stream_info.matches_extension(['.docx']))
    
    def convert(self, file_stream: BinaryIO, stream_info: StreamInfo, **kwargs) -> DocumentConverterResult:
        """Convert DOCX file"""
        logger.info(f"Converting DOCX: {stream_info.filename}")
        
        if not self.python_docx_available:
            raise MissingDependencyException("python-docx is required for DOCX conversion")
        
        try:
            from docx import Document
            
            # Load document
            file_stream.seek(0)
            doc = Document(file_stream)
            
            # Extract content
            markdown = self._extract_content(doc)
            
            # Extract metadata
            metadata = self._extract_docx_metadata(doc, stream_info)
            
            # Extract title
            title = self._extract_docx_title(doc)
            
            return DocumentConverterResult(
                markdown=self._clean_markdown(markdown),
                title=title,
                metadata=metadata
            )
            
        except Exception as e:
            raise FileConversionException(f"DOCX conversion failed: {e}")
    
    def _extract_content(self, doc) -> str:
        """Extract content from DOCX document"""
        markdown = ""
        
        # Process paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            
            # Check paragraph style for headings
            style_name = paragraph.style.name.lower()
            
            if 'heading' in style_name:
                # Extract heading level
                level = self._extract_heading_level(style_name)
                markdown += f"{'#' * level} {text}\n\n"
            elif 'title' in style_name:
                # Document title
                markdown += f"# {text}\n\n"
            else:
                # Regular paragraph
                # Process inline formatting
                formatted_text = self._process_paragraph_formatting(paragraph)
                markdown += f"{formatted_text}\n\n"
        
        # Process tables
        for table in doc.tables:
            table_markdown = self._convert_table(table)
            if table_markdown:
                markdown += table_markdown + "\n\n"
        
        return markdown
    
    def _extract_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name"""
        import re
        
        # Look for numbers in heading style (e.g., "Heading 1", "Heading 2")
        match = re.search(r'heading\s*(\d+)', style_name)
        if match:
            level = int(match.group(1))
            return min(level, 6)  # Max heading level is 6
        
        return 1  # Default to h1
    
    def _process_paragraph_formatting(self, paragraph) -> str:
        """Process inline formatting in paragraph"""
        text = ""
        
        for run in paragraph.runs:
            run_text = run.text
            
            # Apply formatting
            if run.bold and run.italic:
                run_text = f"***{run_text}***"
            elif run.bold:
                run_text = f"**{run_text}**"
            elif run.italic:
                run_text = f"*{run_text}*"
            
            # Handle code formatting (assuming Courier New or similar)
            font_name = run.font.name or ""
            if any(font in font_name.lower() for font in ['courier', 'consola', 'monaco']):
                run_text = f"`{run_text}`"
            
            text += run_text
        
        return text
    
    def _convert_table(self, table) -> str:
        """Convert DOCX table to markdown"""
        rows = []
        
        # Extract table data
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                # Escape pipe characters
                cell_text = cell_text.replace('|', '\\|')
                # Replace newlines with spaces
                cell_text = cell_text.replace('\n', ' ')
                cells.append(cell_text)
            
            if cells:
                rows.append(cells)
        
        if not rows:
            return ""
        
        # Create markdown table
        markdown = ""
        
        # Header row
        if rows:
            header = rows[0]
            markdown += "| " + " | ".join(header) + " |\n"
            markdown += "| " + " | ".join(["---"] * len(header)) + " |\n"
            
            # Data rows
            for row in rows[1:]:
                # Adjust row length to match header
                row = row[:len(header)]
                row.extend([""] * (len(header) - len(row)))
                markdown += "| " + " | ".join(row) + " |\n"
        
        return markdown
    
    def _extract_docx_metadata(self, doc, stream_info: StreamInfo) -> Dict[str, Any]:
        """Extract DOCX metadata"""
        metadata = self._extract_metadata(None, stream_info)
        
        # Document properties
        core_props = doc.core_properties
        
        if core_props.title:
            metadata['title'] = core_props.title
        if core_props.author:
            metadata['author'] = core_props.author
        if core_props.subject:
            metadata['subject'] = core_props.subject
        if core_props.keywords:
            metadata['keywords'] = core_props.keywords
        if core_props.comments:
            metadata['comments'] = core_props.comments
        if core_props.category:
            metadata['category'] = core_props.category
        if core_props.created:
            metadata['created'] = core_props.created.isoformat()
        if core_props.modified:
            metadata['modified'] = core_props.modified.isoformat()
        if core_props.last_modified_by:
            metadata['last_modified_by'] = core_props.last_modified_by
        if core_props.revision:
            metadata['revision'] = core_props.revision
        
        # Document statistics
        try:
            # Count paragraphs and tables
            metadata['paragraph_count'] = len(doc.paragraphs)
            metadata['table_count'] = len(doc.tables)
            
            # Count images (shapes with image data)
            image_count = 0
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if hasattr(run, '_element') and run._element.xpath('.//a:blip'):
                        image_count += 1
            metadata['image_count'] = image_count
            
        except Exception as e:
            logger.debug(f"Could not extract document statistics: {e}")
        
        return metadata
    
    def _extract_docx_title(self, doc) -> Optional[str]:
        """Extract title from DOCX document"""
        # Try document properties first
        if doc.core_properties.title:
            return self._format_title(doc.core_properties.title)
        
        # Try first heading or title-styled paragraph
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            
            style_name = paragraph.style.name.lower()
            
            if 'title' in style_name or 'heading 1' in style_name:
                return self._format_title(text)
        
        # Try first non-empty paragraph
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text and len(text) < 100:
                return self._format_title(text)
        
        return None
    
    def get_format_info(self) -> Dict[str, Any]:
        """Get format information"""
        return {
            'name': 'DOCX',
            'description': 'Microsoft Word document converter with formatting preservation',
            'extensions': self.supported_extensions,
            'mimetypes': self.supported_mimetypes,
            'category': self.category,
            'features': [
                'Text extraction with formatting',
                'Heading detection',
                'Table conversion',
                'Metadata extraction',
                'Inline formatting (bold, italic)',
                'Korean text support'
            ]
        }